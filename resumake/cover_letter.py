"""Cover letter command — generate a cover letter for a job description."""

from pathlib import Path
from typing import Annotated, Optional

import typer
import yaml
from docx import Document
from docx.shared import Cm, Pt

from .console import console, err_console
from .llm import get_provider, strip_yaml_fences
from .theme import Theme, load_theme
from .utils import DEFAULT_YAML, OUTPUT_DIR, convert_to_pdf, load_cv, open_file, slugify_name


def _generate_cover_letter(cv: dict, description_text: str) -> dict:
    """Use an LLM to generate a cover letter matching the CV to a job description."""
    provider = get_provider()
    cv_yaml = yaml.dump(cv, allow_unicode=True, default_flow_style=False, sort_keys=False)

    with console.status("Generating cover letter via LLM..."):
        response = provider.complete(
            "You are a professional career consultant. Given a CV in YAML format and a job description, "
            "write a compelling cover letter.\n\n"
            "Return ONLY valid YAML with this exact structure — no explanation, no code fences:\n\n"
            "recipient: <company name or 'Hiring Manager'>\n"
            "subject: <subject line for the letter>\n"
            "opening: <opening paragraph — why you're writing and the role you're applying for>\n"
            "body: <1-2 paragraphs connecting your experience to the role requirements>\n"
            "closing: <closing paragraph — call to action, availability, enthusiasm>\n\n"
            "Rules:\n"
            "- Write in first person, professional but personable tone.\n"
            "- Reference specific achievements from the CV that match the job requirements.\n"
            "- Do NOT invent experience or skills not in the CV.\n"
            "- Keep it concise — no more than one page when formatted.\n"
            "- Do not include the sender's address or date — those will be added from CV data.\n\n"
            f"JOB DESCRIPTION:\n{description_text}\n\n"
            f"CV YAML:\n{cv_yaml}",
            max_tokens=2048,
        )

    return yaml.safe_load(strip_yaml_fences(response))


def _build_cover_letter_docx(cv: dict, letter: dict, lang: str, theme: Theme) -> Path:
    """Build a cover letter Word document."""
    t = theme
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = t.fonts.body
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Sender info
    contact = cv.get("contact", {})
    sender_lines = [cv["name"]]
    if contact.get("address"):
        sender_lines.append(contact["address"])
    if contact.get("email"):
        sender_lines.append(contact["email"])
    if contact.get("phone"):
        sender_lines.append(contact["phone"])

    for line in sender_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = t.colors.text_muted_rgb
        run.font.name = t.fonts.heading
        p.paragraph_format.space_after = Pt(1)

    # Spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Recipient
    recipient = letter.get("recipient", "Hiring Manager")
    p = doc.add_paragraph()
    run = p.add_run(f"Dear {recipient},")
    run.font.size = Pt(11)
    run.font.name = t.fonts.body
    run.bold = True
    p.paragraph_format.space_after = Pt(12)

    # Subject line
    if letter.get("subject"):
        p = doc.add_paragraph()
        run = p.add_run(f"Re: {letter['subject']}")
        run.font.size = Pt(11)
        run.font.name = t.fonts.heading
        run.bold = True
        run.font.color.rgb = t.colors.primary_rgb
        p.paragraph_format.space_after = Pt(12)

    # Body paragraphs
    for key in ("opening", "body", "closing"):
        text = letter.get(key, "")
        if text:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.name = t.fonts.body
            run.font.color.rgb = t.colors.text_body_rgb
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = 1.15

    # Sign-off
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    run = p.add_run("Sincerely,")
    run.font.size = Pt(11)
    run.font.name = t.fonts.body
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run(cv["name"])
    run.font.size = Pt(11)
    run.font.name = t.fonts.heading
    run.bold = True
    run.font.color.rgb = t.colors.primary_rgb

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"{slugify_name(cv['name'])}_Cover_Letter_{lang.upper()}.docx"
    output_path = OUTPUT_DIR / filename
    doc.save(str(output_path))
    return output_path


def cover_letter(
    description_file: Annotated[
        Path,
        typer.Argument(help="Path to a .txt or .md file with the job description."),
    ],
    lang: Annotated[Optional[str], typer.Option(help="Output language (en or de). Default: en.")] = "en",
    source: Annotated[Path, typer.Option(help="Path to source YAML.")] = DEFAULT_YAML,
    pdf: Annotated[bool, typer.Option("--pdf", help="Also generate PDF.")] = False,
    open: Annotated[bool, typer.Option("--open/--no-open", help="Open the generated file.")] = True,
    theme: Annotated[
        Optional[str],
        typer.Option(help="Theme name or path to theme.yaml."),
    ] = None,
):
    """Generate a cover letter for a job description based on your CV."""
    if not description_file.exists():
        err_console.print(f"[red]Error:[/] File not found: {description_file}")
        raise typer.Exit(1)

    description_text = description_file.read_text(encoding="utf-8").strip()
    if not description_text:
        err_console.print("[red]Error:[/] Description file is empty.")
        raise typer.Exit(1)

    cv = load_cv(source)
    letter = _generate_cover_letter(cv, description_text)

    resolved_theme = load_theme(theme)
    output_path = _build_cover_letter_docx(cv, letter, lang, resolved_theme)
    console.print(f"Generated: [cyan]{output_path}[/]")

    outputs = [output_path]
    if pdf:
        pdf_path = convert_to_pdf(output_path)
        console.print(f"Generated: [cyan]{pdf_path}[/]")
        outputs.append(pdf_path)

    if open:
        for p in outputs:
            open_file(p)
