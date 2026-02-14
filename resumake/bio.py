"""Bio command — generate a condensed one-pager bio."""

from pathlib import Path
from typing import Annotated, Optional

import typer
import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt

from .docx_builder import (
    add_hyperlink,
    remove_table_borders,
    set_cell_width,
)
from .theme import Theme, load_theme
from .translate import translate_cv
from .utils import (
    DEFAULT_YAML,
    OUTPUT_DIR,
    convert_to_pdf,
    load_cv,
    open_file,
    parse_start_date,
    resolve_asset,
    slugify_name,
)


def select_bio_content(cv: dict) -> dict:
    """Use an LLM to select and condense CV content for a one-pager bio."""
    from .llm import get_provider, strip_yaml_fences

    provider = get_provider()

    cv_yaml = yaml.dump(cv, allow_unicode=True, default_flow_style=False, sort_keys=False)

    from .console import console

    with console.status("Generating bio content via LLM..."):
        response = provider.complete(
            "You are a professional CV consultant. Given a full CV in YAML, produce a condensed "
            "one-pager bio version.\n\n"
            "Return ONLY valid YAML with this exact structure — no explanation, no code fences:\n\n"
            "name: <full name>\n"
            "title: <professional title>\n"
            "photo: <photo path from original>\n"
            "contact:\n"
            "  email: <email>\n"
            "  phone: <phone>\n"
            "  address: <city, country>\n"
            "bio_summary: <3-4 sentences in third person summarizing the person's career, "
            "expertise, and value proposition>\n"
            "career_highlights:\n"
            "  - <highlight 1>\n"
            "  - <highlight 2>\n"
            "  - <highlight 3>\n"
            "  - <highlight 4>\n"
            "  - <highlight 5>\n"
            "current_roles:\n"
            "  - title: <title>\n"
            "    org: <org>\n"
            "    period: <start — end>\n"
            "  ...(2-3 most recent roles)\n"
            "education:\n"
            "  - <degree, institution>\n"
            "  ...\n"
            "skills_summary: <comma-separated list of 10-15 key skills>\n"
            "links:\n"
            "  - label: <label>\n"
            "    url: <url>\n"
            "  ...\n\n"
            "Rules:\n"
            "- Do NOT invent content. Only select and condense from the original CV.\n"
            "- Career highlights should be the most impressive, quantifiable achievements.\n"
            "- Current roles = 2-3 most recent experience entries.\n"
            "- Skills summary = most important skills across all categories.\n"
            "- Bio summary should be written in third person, professional tone.\n\n"
            f"CV YAML:\n{cv_yaml}",
        )

    return yaml.safe_load(strip_yaml_fences(response))


def select_bio_content_deterministic(cv: dict) -> dict:
    """Deterministic fallback: select bio content without using the API."""
    # Sort experience by date, take top 3
    entries = sorted(
        cv.get("experience", []),
        key=lambda e: parse_start_date(e.get("start", "")),
        reverse=True,
    )
    current_roles = [
        {"title": e["title"], "org": e.get("org", ""), "period": f"{e['start']} — {e['end']}"} for e in entries[:3]
    ]

    # Collect top bullets as highlights
    highlights = []
    for e in entries[:4]:
        for b in e.get("bullets", [])[:2]:
            highlights.append(b)
            if len(highlights) >= 5:
                break
        if len(highlights) >= 5:
            break

    # Skills summary
    skills = cv.get("skills", {})
    all_skills = skills.get("leadership", []) + skills.get("technical", [])
    skills_summary = ", ".join(all_skills[:15])

    # Education
    education = [f"{e['degree']}, {e['institution']}" for e in cv.get("education", [])]

    # Profile as bio summary
    bio_summary = cv.get("profile", "").strip()
    # Take first 3 sentences
    sentences = bio_summary.split(". ")
    bio_summary = ". ".join(sentences[:4])
    if not bio_summary.endswith("."):
        bio_summary += "."

    return {
        "name": cv["name"],
        "title": cv["title"],
        "photo": cv.get("photo", ""),
        "contact": {
            "email": cv["contact"]["email"],
            "phone": cv["contact"]["phone"],
            "address": cv["contact"]["address"],
        },
        "bio_summary": bio_summary,
        "career_highlights": highlights,
        "current_roles": current_roles,
        "education": education,
        "skills_summary": skills_summary,
        "links": cv.get("links", []),
    }


def build_bio_docx(bio_data: dict, lang: str, theme: Theme | None = None) -> Path:
    """Build a single-column one-pager bio document."""
    t = theme or load_theme()
    doc = Document()

    # Page setup — wider margins for single column
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = t.fonts.heading
    style.font.size = Pt(t.sizes.body_pt)
    style.font.color.rgb = t.colors.text_body_rgb
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ── Header: photo + name + title + contact in a table ──
    header_table = doc.add_table(rows=1, cols=2)
    remove_table_borders(header_table)

    left_cell = header_table.cell(0, 0)
    right_cell = header_table.cell(0, 1)

    # Set widths
    set_cell_width(left_cell, 4.0)
    set_cell_width(right_cell, 13.0)

    # Photo in left cell
    photo_path = resolve_asset(bio_data["photo"]) if bio_data.get("photo") else None
    if photo_path and photo_path.exists():
        p_photo = left_cell.paragraphs[0]
        p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_photo.add_run()
        run.add_picture(str(photo_path), width=Cm(3.0))

    # Name + title + contact in right cell
    p_name = right_cell.paragraphs[0]
    run_name = p_name.add_run(bio_data["name"])
    run_name.bold = True
    run_name.font.size = Pt(t.sizes.name_pt + 5)
    run_name.font.color.rgb = t.colors.primary_rgb
    run_name.font.name = t.fonts.heading

    p_title = right_cell.add_paragraph()
    run_title = p_title.add_run(bio_data["title"])
    run_title.font.size = Pt(t.sizes.subheading_pt + 1)
    run_title.font.color.rgb = t.colors.accent_rgb
    run_title.font.name = t.fonts.heading
    run_title.italic = True
    p_title.paragraph_format.space_after = Pt(6)

    # Contact line
    contact = bio_data.get("contact", {})
    contact_parts = [v for v in [contact.get("email"), contact.get("phone"), contact.get("address")] if v]
    if contact_parts:
        p_contact = right_cell.add_paragraph()
        run_c = p_contact.add_run("  |  ".join(contact_parts))
        run_c.font.size = Pt(t.sizes.small_pt)
        run_c.font.color.rgb = t.colors.text_muted_rgb
        run_c.font.name = t.fonts.heading

    # Links
    links = bio_data.get("links", [])
    if links:
        p_links = right_cell.add_paragraph()
        for i, link in enumerate(links):
            if i > 0:
                sep = p_links.add_run("  |  ")
                sep.font.size = Pt(t.sizes.small_pt)
                sep.font.color.rgb = t.colors.text_muted_rgb
                sep.font.name = t.fonts.heading
            if link.get("url"):
                add_hyperlink(
                    p_links,
                    link["url"],
                    link["label"],
                    color=t.colors.accent_rgb,
                    size=Pt(t.sizes.small_pt),
                    font_name=t.fonts.heading,
                )
            else:
                run_l = p_links.add_run(link["label"])
                run_l.font.size = Pt(t.sizes.small_pt)
                run_l.font.color.rgb = t.colors.accent_rgb
                run_l.font.name = t.fonts.heading

    # ── Horizontal divider ──
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(8)
    p_div.paragraph_format.space_after = Pt(8)
    pPr = p_div._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{t.colors.accent}"/></w:pBdr>'
    )
    pPr.append(pBdr)

    # ── Bio summary ──
    def add_section_title(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(t.sizes.subheading_pt + 1)
        run.font.color.rgb = t.colors.primary_rgb
        run.font.name = t.fonts.heading

    add_section_title("Professional Summary")
    p_bio = doc.add_paragraph()
    run_bio = p_bio.add_run(bio_data.get("bio_summary", ""))
    run_bio.font.size = Pt(t.sizes.body_pt + 1)
    run_bio.font.color.rgb = t.colors.text_body_rgb
    run_bio.font.name = t.fonts.body
    p_bio.paragraph_format.space_after = Pt(6)

    # ── Career highlights ──
    highlights = bio_data.get("career_highlights", [])
    if highlights:
        add_section_title("Key Achievements")
        for h in highlights:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.text = ""
            run = p.add_run(h)
            run.font.size = Pt(t.sizes.body_pt)
            run.font.color.rgb = t.colors.text_body_rgb
            run.font.name = t.fonts.body

    # ── Current roles ──
    roles = bio_data.get("current_roles", [])
    if roles:
        add_section_title("Recent Experience")
        for role in roles:
            p_role = doc.add_paragraph()
            p_role.paragraph_format.space_before = Pt(4)
            p_role.paragraph_format.space_after = Pt(1)
            run_t = p_role.add_run(f"{role['title']} — {role.get('org', '')}")
            run_t.bold = True
            run_t.font.size = Pt(t.sizes.body_pt)
            run_t.font.color.rgb = t.colors.primary_rgb
            run_t.font.name = t.fonts.heading

            p_period = doc.add_paragraph()
            run_p = p_period.add_run(role.get("period", ""))
            run_p.font.size = Pt(t.sizes.small_pt)
            run_p.font.color.rgb = t.colors.text_muted_rgb
            run_p.font.name = t.fonts.heading

    # ── Education ──
    education = bio_data.get("education", [])
    if education:
        add_section_title("Education")
        for edu in education:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_after = Pt(2)
            edu_text = edu if isinstance(edu, str) else f"{edu.get('degree', '')}, {edu.get('institution', '')}"
            run_edu = p_edu.add_run(edu_text)
            run_edu.font.size = Pt(t.sizes.body_pt)
            run_edu.font.color.rgb = t.colors.text_body_rgb
            run_edu.font.name = t.fonts.heading

    # ── Skills summary ──
    skills_summary = bio_data.get("skills_summary", "")
    if skills_summary:
        add_section_title("Core Competencies")
        p_skills = doc.add_paragraph()
        run_skills = p_skills.add_run(skills_summary)
        run_skills.font.size = Pt(t.sizes.body_pt)
        run_skills.font.color.rgb = t.colors.text_body_rgb
        run_skills.font.name = t.fonts.body

    # Save
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"{slugify_name(bio_data['name'])}_Bio_{lang.upper()}.docx"
    output_path = OUTPUT_DIR / filename
    doc.save(str(output_path))
    return output_path


def bio(
    lang: Annotated[Optional[str], typer.Option(help="Output language code (e.g. en, de, fr). Default: en.")] = "en",
    source: Annotated[Path, typer.Option(help="Path to source YAML.")] = DEFAULT_YAML,
    pdf: Annotated[bool, typer.Option("--pdf", help="Also generate PDF.")] = False,
    open: Annotated[bool, typer.Option("--open/--no-open", help="Open the generated file.")] = True,
    theme: Annotated[
        Optional[str],
        typer.Option(help="Theme name (classic, minimal, modern) or path to theme.yaml."),
    ] = None,
):
    """Generate a condensed one-pager bio document."""
    cv_en = load_cv(source)

    from .console import console

    # Try LLM, fall back to deterministic
    try:
        from .llm import get_provider

        get_provider()  # check availability
        bio_data = select_bio_content(cv_en)
    except RuntimeError:
        console.print("[yellow]No LLM provider available — using deterministic bio selection.[/]")
        bio_data = select_bio_content_deterministic(cv_en)

    # Translate bio data if needed
    if lang != "en":
        bio_data = translate_cv(bio_data, lang=lang, retranslate=True)

    resolved_theme = load_theme(theme)
    output_path = build_bio_docx(bio_data, lang, theme=resolved_theme)
    console.print(f"Generated: [cyan]{output_path}[/]")

    outputs = [output_path]
    if pdf:
        pdf_path = convert_to_pdf(output_path)
        console.print(f"Generated: [cyan]{pdf_path}[/]")
        outputs.append(pdf_path)

    if open:
        for p in outputs:
            open_file(p)
